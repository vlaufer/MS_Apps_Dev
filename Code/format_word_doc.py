from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.dml.color import ColorFormat
from docx.shared import RGBColor


picture_path="/Users/vincentlaufer/Desktop/Consultancies/Laufer_D/Automated_Excel_Functions/Results/Resized_Pictures/IBM Tabulator.jpg"

## first, initialize your document object:
document = Document()




# p.add_run('Document Title', 0).bold = True
# p = document.add_paragraph()
# p.add_run('Doc text').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

document.add_picture(picture_path, width=Inches(4))


# Construct object with dictionary
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('/Users/vincentlaufer/Desktop/demo.docx')

#need to figure out:
	# bold / italic
	# justification
	# text highlight
	# text color		--> docx.text.run.Font[source] color
	# size






