##################
# import modules #
##################

# for processing picture files and word documents
from PIL import Image
import PIL.ExifTags

# for file I/O
import glob, os

# for functionality necessary to write to MS Word.
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

print("\n\n\nThis software is distributed under the MIT license, please see https://opensource.org/licenses/MIT for more information.")
print("Please contact vincent.a.laufer@gmail.com with questions, bugs, etc.\n\n--- starting processing ---\n")

def obtain_header_indices(colname_array, source_tsv):
## this function indentifies the row containing the header, and returns the row number that contains it as an integer.
	print("Parsing the input Excel file ... \n")
	header_candidate_list=[]; header_length_list=[]; header_dict=dict()
	i=0; j=0
	source_text_dictionary=dict()

	with open(source_tsv, 'r') as input_stream:
		header=input_stream.readline()
		header=header.strip().split("\t")
		header_proc=[ field.lower() for field in header ]
		for colname in colname_array:
			header_dict[colname]=header_proc.index(colname)
		for line  in input_stream:
			line=line.strip().split('\t')
			line[4]=line[4].strip()
			source_text_dictionary[ line[3] ] = line
	return(source_text_dictionary)


def create_integrated_Word_document(input_text_dictionary, picture_dir, result_dir):
	print("Writing the curated text and pictures to a custom MS Word Document ... \n")
	document = Document() ## first, initialize your document object:
	i=0


	styles = document.styles
	header_custom_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
	header_custom_style.font.name = 'Arial'
	header_custom_style.font.size = Pt(24)
	header_custom_style.font.bold = True

	body_custom_style = styles.add_style('Section Text', WD_STYLE_TYPE.PARAGRAPH)
	body_custom_style.font.name = 'Arial'
	body_custom_style.font.size = Pt(12)
	body_custom_style.font.bold = True
	body_custom_style.font.highlight_color = WD_COLOR_INDEX.WHITE

	for milestone_name in input_text_dictionary.keys():

# make new section for each record.

# define input variables:
		category=input_text_dictionary[milestone_name][0]
		start_date=input_text_dictionary[milestone_name][1]

		if input_text_dictionary[milestone_name][2] == "-":
			end_date="\t"
		else: 
			end_date=input_text_dictionary[milestone_name][2]

		milestone_text=input_text_dictionary[milestone_name][4]

		if input_text_dictionary[milestone_name][5] == "-":
			picture_path="No picture provided"
		else:
			picture_path=picture_dir + "/" + input_text_dictionary[milestone_name][5]

##########################################################
# Define your user custom style for the header #
##########################################################

# Assign the theme color based on the category from column A of the excel sheet.
		if category == 'MACHINES':
			header_custom_style.font.color.rgb = RGBColor(250, 250, 250)
			header_custom_style.font.highlight_color = WD_COLOR_INDEX.BLUE
		if category == 'INNOVATIONS':
			header_custom_style.font.color.rgb = RGBColor(250, 250, 250)
			header_custom_style.font.highlight_color = WD_COLOR_INDEX.GREEN
		if category == 'PEOPLE':
			header_custom_style.font.color.rgb = RGBColor(250, 250, 250)
			header_custom_style.font.highlight_color = WD_COLOR_INDEX.RED

######################################
# Now actually add the header #
######################################
		section_header=document.add_paragraph(milestone_name, style=header_custom_style) # Now actually add the header to the MS Word file. 

########################
# Now insert the photo #
########################
		if picture_path == "No picture provided":
			document.add_paragraph("No Picture provided for this vignette") 
		else:
			document.add_picture(picture_path, width=Inches(4)) # add picture

###################################################
# Define your user custom style for the body text #
###################################################
# Assign the theme color based on the category from column A of the excel sheet.
		if category == 'MACHINES':
			body_custom_style.font.color.rgb = RGBColor(0, 0, 0)
		if category == 'INNOVATIONS':
			body_custom_style.font.color.rgb = RGBColor(0, 0, 0)
		if category == 'PEOPLE':
			body_custom_style.font.color.rgb = RGBColor(0, 0, 0)

		info_paragraph=document.add_paragraph( style = body_custom_style)
		if end_date != "\t":
			info_paragraph.add_run(start_date + " - " + end_date + ": ").bold=True # Now actually add the header to the MS Word file. 
		else:
			info_paragraph.add_run(start_date + ": ").bold=True
		info_paragraph.add_run(milestone_text).bold = False # Now actually add the header to the MS Word file. 

		document.add_page_break() # since each record gets its own page, end the page once the section milestone text is complete.
		i+=1

	document.save(result_dir + '/CS_museum_mockup.docx')

#need to figure out:
	# justification 		
	# text highlight 		
	# text color			docx.text.run.Font[source] color
	# size 					


def main():

###################################################
# Define input variables based on client desires. #
###################################################
	column_name_array_raw=["CATEGORY", "START DATE", "END DATE", "MILESTONE NAME", "MILESTONE TEXT", "image name"] 	# Take as input the names of the columns corresponding to each field that you wish to capture for a given record.
	column_name_array=[ field.lower() for field in column_name_array_raw ]

# input file paths:
	base_dir="/Users/vincentlaufer/Desktop/Consultancies/Laufer_D/Automated_Excel_Functions"
	script_dir=base_dir + "/" + "Code"
	data_dir=base_dir + "/" + "Data"
	picture_dir=data_dir + "/" +"Source_Pictures"
	output_dir=base_dir + "/" + "Results"

# individual file names:
	text_source_file= data_dir + "/" "sample_input_reformatted.txt"

#################################################
# First, parse the input text into a dictionary #
#################################################
	text_dictionary=obtain_header_indices(column_name_array, text_source_file)

######################################################
# Second, input the pictures and the text into word. #
######################################################
	create_integrated_Word_document( text_dictionary, picture_dir, output_dir)
	print("Processing completed successfully. Output MS Word file is in " + output_dir + ".\n")

if __name__ == "__main__": 
	main()










