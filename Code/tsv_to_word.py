# import modules we will need to process picture files and word documents
from PIL import Image
import PIL.ExifTags

# import modules needed for file I/O
import glob, os

# import functionality necessary to write to MS Word.
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.dml.color import ColorFormat
from docx.shared import RGBColor



print("\n\n\nThis software is distributed under the MIT license, please see https://opensource.org/licenses/MIT for more information.")
print("Please contact vincent.a.laufer@gmail.com with questions, bugs, etc.\n\n--- starting processing ---\n")

def obtain_header_indices(colname_array, source_tsv):
## this function indentifies the row containing the header, and returns the row number that contains it as an integer.
	print("Parsing the input Excel file ... \n")
	header_candidate_list=[]; header_length_list=[]; header_dict=dict()
	i=0; j=0
	row_match_index=-1
	source_text_dictionary=dict()

	with open(source_tsv, 'r') as input_stream:
		while row_match_index == -1:
			header_candidate=input_stream.readline()
			header_candidate=header_candidate.strip().split("\t")
			header_candidate_proc=[ field.lower() for field in header_candidate ]
			num_field_matches=set(header_candidate_proc) & set(colname_array)
			if len(num_field_matches) == len(colname_array):
				row_match_index=i
				header=header_candidate_proc
				for colname in colname_array:
					header_dict[colname]=header.index(colname)
			i+=1
		if row_match_index != -1:
			for line  in input_stream:
				line=line.strip().split('\t')
				line[4]=line[4].strip()
				source_text_dictionary[ line[3] ] = line
	return(source_text_dictionary)


#####################################################################################################################################
############ Important - maybe possible to totally remove this function and do it all through docx. Need to investigate. ############
#####################################################################################################################################
def standardize_picture_files(picture_input_dir, target_width_inches, picture_file_type, picture_output_dir):
	print("Standardizing the input picture files ... \n")
	input_picture_paths=[]; output_picture_paths=[]
	for root, dirs, files in os.walk(picture_input_dir):
		for file in files:
## file naming:
			input_picture_fname=os.path.join(picture_input_dir, file)
			output_picture_fname=os.path.join(picture_output_dir, file)
## calculation of original dimensions
			if os.path.isfile(output_picture_fname) == False:
				try:
					current_image=Image.open(input_picture_fname) ## open the source file
					pixel_width, pixel_height = current_image.size  ## get the number of pixels and the dpi of the original image.
					print(pixel_width); print(pixel_height)
					dimensional_ratio=(pixel_width / 288)
					final_width_in_pixels=288
					final_height_in_pixels=(pixel_height / dimensional_ratio)
					new_pixel_dimensions=( int(final_width_in_pixels), int(final_height_in_pixels)) ## now get the new size in pixels, using the same ratio for length and width to preserve aspect ratio
					processed_image=current_image.resize(new_pixel_dimensions) 
					processed_image.save(output_picture_fname) ## processed file output
				except PIL.UnidentifiedImageError:
					print("Warning: An unidentified file has been detected in the picture folder.\nAdding " + input_picture_fname + " to the logfile for tracking, and skipping this file...\n")
#####################################################################################################################################
############ Important - maybe possible to totally remove this function and do it all through docx. Need to investigate. ############
#####################################################################################################################################

def create_integrated_Word_document(input_text_dictionary, picture_dir, result_dir):
	print("Writing the curated text and pictures to a custom MS Word Document ... \n")
	document = Document() ## first, initialize your document object:
	i=0

	for milestone_name in input_text_dictionary.keys():

# make new section for each record.
		section = document.sections[i]

# define input variables:
		category=input_text_dictionary[milestone_name][0]
		start_date=input_text_dictionary[milestone_name][1]
		end_date=input_text_dictionary[milestone_name][2]
		milestone_text=input_text_dictionary[milestone_name][3]
		image_name=input_text_dictionary[milestone_name][4]
		picture_path=picture_dir + "/" + input_text_dictionary[milestone_name][5]

#########################
# Make the header style #
#########################
		header_style = document.styles['Normal']
		header_font = header_style.font
		header_font.name = 'Arial'
		header_font.size = Pt(16)
		header_font.bold = True

# Assign the theme color based on the category from column A of the excel sheet.
		if category == 'MACHINES':
			header_font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)
		if category == 'INNOVATIONS':
			header_font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
		if category == 'PEOPLE':
			header_font.color.rgb = RGBColor(0x36, 0x24, 0x3f)

######################################
# Now actually add the styled header #
######################################
		section_header=document.add_paragraph(milestone_name, style = header_style) # Now actually add the header to the MS Word file. 


########################
# Now insert the photo #
########################
		document.add_picture(picture_path, width=Inches(4)) # add picture

#################################
# Make the milestone text style #
#################################
		paragraph_text_style = document.styles['Normal']
		paragraph_text_font = paragraph_text_style.font
		paragraph_text_font.name = 'Arial'
		paragraph_text_font.size = Pt(12)
		paragraph_text_font.bold = True

# Assign the theme color based on the category from column A of the excel sheet.
		if category == 'MACHINES':
			paragraph_text_font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)
		if category == 'INNOVATIONS':
			paragraph_text_font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
		if category == 'PEOPLE':
			paragraph_text_font.color.rgb = RGBColor(0x36, 0x24, 0x3f)

		section_info=document.add_paragraph(start_date + " - " + end_date + ": ", style = paragraph_text_style) # Now actually add the header to the MS Word file. 
		paragraph_text_font.bold = False # turn bold off now.
		section_info=document.add_run(milestone_text, style = paragraph_text_style) # Now actually add the header to the MS Word file. 

		document.add_page_break() # since each record gets its own page, end the page once the section milestone text is complete.
		i+=1

	document.save('/Users/vincentlaufer/Desktop/demo.docx') #after everything, 




# 

# save file
	document.save(result_dir + '/CS_museum_mockup.docx')
#need to figure out:
	# bold / italic
	# justification
	# text highlight
	# text color		--> docx.text.run.Font[source] color
	# size


def main():

###################################################
# Define input variables based on client desires. #
###################################################
	column_name_array_raw=["CATEGORY", "START DATE", "END DATE", "MILESTONE NAME", "MILESTONE TEXT", "image name"] 	# Take as input the names of the columns corresponding to each field that you wish to capture for a given record.
	column_name_array=[ field.lower() for field in column_name_array_raw ]

# Picture file parameters
	picture_width_inches=4
	picture_file_type="jpg"
	compression_format="lzw"

# Font specifications (body)
	header_font_size=18				# 
	header_font_name="Arial" 			# Arial Narrow will allow more characters per line.
	header_line_spacing=1.15			# 

# Font specifications (header)
	body_font_size=11
	body_font_name="Arial"
	body_line_spacing=1.15

# input file paths:
	base_dir="/Users/vincentlaufer/Desktop/Consultancies/Laufer_D/Automated_Excel_Functions"
	script_dir=base_dir + "/" + "Code"
	data_dir=base_dir + "/" +"Data"
	picture_dir=data_dir + "/" +"Source_Pictures"
	output_dir=base_dir + "/" + "Results"
	picture_output_dir=output_dir + "/Resized_Pictures" 

# individual file names:
	text_source_file= data_dir + "/" "SC_timeline.tsv"

#################################################
# First, parse the input text into a dictionary #
#################################################
	text_dictionary=obtain_header_indices(column_name_array, text_source_file)

#################################################
# Second, standardize all pictures to 4" .jpegs #
#################################################
	standardize_picture_files(picture_dir, picture_width_inches, picture_file_type, picture_output_dir)

#####################################################
# Third, input the pictures and the text into word. #
#####################################################
	create_integrated_Word_document( text_dictionary, picture_output_dir, output_dir)
	print("Processing completed successfully. Output MS Word file is in " + output_dir + ".\n")

if __name__ == "__main__": 
	main()










